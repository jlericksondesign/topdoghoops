from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "privacy" / "review-docs"

SOURCES = [
    (
        "Top Dog Hoops Privacy Policy",
        "Privacy Policy Draft",
        ROOT / "docs" / "privacy" / "privacy-policy-draft.md",
        OUT_DIR / "top-dog-hoops-privacy-policy-review.docx",
    ),
    (
        "Top Dog Hoops Terms of Use",
        "Terms of Use Draft",
        ROOT / "docs" / "privacy" / "terms-of-use-draft.md",
        OUT_DIR / "top-dog-hoops-terms-of-use-review.docx",
    ),
    (
        "Top Dog Hoops Parent Consent and Deletion Request Flow",
        "Parent Consent and Deletion Request Flow Draft",
        ROOT / "docs" / "privacy" / "coppa-parent-consent-flow.md",
        OUT_DIR / "top-dog-hoops-parent-consent-deletion-review.docx",
    ),
]


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_run_font(run, *, size=11, bold=False, italic=False, color="000000"):
    font = run.font
    font.name = "Arial"
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
      rfonts = OxmlElement("w:rFonts")
      rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")


def configure_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, "000000"),
        ("Heading 2", 16, 18, 6, "000000"),
        ("Heading 3", 14, 16, 4, "434343"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def create_numbering_instance(doc: Document):
    numbering = doc.part.numbering_part.element
    decimal_abstract_id = None

    for abstract_num in numbering.findall(qn("w:abstractNum")):
        for num_fmt in abstract_num.findall(".//" + qn("w:numFmt")):
            if num_fmt.get(qn("w:val")) == "decimal":
                decimal_abstract_id = abstract_num.get(qn("w:abstractNumId"))
                break
        if decimal_abstract_id is not None:
            break

    if decimal_abstract_id is None:
        decimal_abstract_id = "0"

    existing_ids = [
        int(num.get(qn("w:numId")))
        for num in numbering.findall(qn("w:num"))
        if num.get(qn("w:numId")) and num.get(qn("w:numId")).isdigit()
    ]
    next_num_id = str(max(existing_ids, default=0) + 1)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), next_num_id)
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), decimal_abstract_id)
    num.append(abstract_num_id)

    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)

    return next_num_id


def apply_numbering(paragraph, num_id: str):
    ppr = paragraph._p.get_or_add_pPr()
    for existing in ppr.findall(qn("w:numPr")):
        ppr.remove(existing)

    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), num_id)
    num_pr.append(ilvl)
    num_pr.append(num_id_element)
    ppr.append(num_pr)


def add_title(doc: Document, title: str, subtitle: str):
    title_p = doc.add_paragraph()
    set_paragraph_spacing(title_p, before=0, after=3, line=1.15)
    run = title_p.add_run(title)
    set_run_font(run, size=26, bold=False)

    subtitle_p = doc.add_paragraph()
    set_paragraph_spacing(subtitle_p, before=0, after=12, line=1.15)
    run = subtitle_p.add_run(subtitle)
    set_run_font(run, size=11, italic=True, color="555555")

    note = doc.add_paragraph()
    set_paragraph_spacing(note, before=0, after=14, line=1.15)
    run = note.add_run(
        "Review mode: please use Google Docs comments or suggestions for proposed edits. "
        "Bracketed placeholders mark information that must be completed before launch."
    )
    set_run_font(run, size=11, italic=True, color="555555")


INLINE_CODE = re.compile(r"`([^`]+)`")
BOLD = re.compile(r"\*\*([^*]+)\*\*")


def add_runs_with_markup(paragraph, text: str):
    cursor = 0
    for match in re.finditer(r"`([^`]+)`|\*\*([^*]+)\*\*", text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run)
        value = match.group(1) or match.group(2) or ""
        run = paragraph.add_run(value)
        set_run_font(run, bold=bool(match.group(2)))
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run)


def add_markdown_content(doc: Document, markdown: str):
    lines = markdown.splitlines()
    i = 0
    active_num_id = None
    while i < len(lines):
        raw = lines[i].rstrip()
        line = raw.strip()
        i += 1

        if not line:
            active_num_id = None
            continue
        if line.startswith("# "):
            active_num_id = None
            continue
        if line.startswith("## "):
            active_num_id = None
            p = doc.add_paragraph(line[3:].strip(), style="Heading 1")
            set_paragraph_spacing(p, before=20, after=6, line=1.15)
            continue
        if line.startswith("### "):
            active_num_id = None
            p = doc.add_paragraph(line[4:].strip(), style="Heading 2")
            set_paragraph_spacing(p, before=18, after=6, line=1.15)
            continue
        if line.startswith("> "):
            active_num_id = None
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            set_paragraph_spacing(p, before=4, after=8, line=1.15)
            run = p.add_run(line[2:].strip())
            set_run_font(run, italic=True)
            continue
        bullet_match = re.match(r"^\* (.+)$", line)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, before=0, after=4, line=1.15)
            add_runs_with_markup(p, bullet_match.group(1).strip())
            continue
        number_match = re.match(r"^\d+\. (.+)$", line)
        if number_match:
            if active_num_id is None:
                active_num_id = create_numbering_instance(doc)
            p = doc.add_paragraph()
            apply_numbering(p, active_num_id)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            set_paragraph_spacing(p, before=0, after=4, line=1.15)
            add_runs_with_markup(p, number_match.group(1).strip())
            continue

        active_num_id = None
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=8, line=1.15)
        add_runs_with_markup(p, line)


def add_review_section(doc: Document):
    doc.add_paragraph("Reviewer Notes", style="Heading 1")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=8, line=1.15)
    run = p.add_run("Use this space for summary comments, approval notes, or required edits.")
    set_run_font(run, italic=True, color="555555")

    for label in [
        "Reviewer name:",
        "Reviewer role:",
        "Approved / approved with edits / needs follow-up:",
        "Summary notes:",
    ]:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=8, line=1.15)
        run = p.add_run(label)
        set_run_font(run, bold=True)


def build_doc(title: str, subtitle: str, source: Path, output: Path):
    doc = Document()
    configure_styles(doc)
    add_title(doc, title, subtitle)
    add_markdown_content(doc, source.read_text())
    add_review_section(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    for title, subtitle, source, output in SOURCES:
        build_doc(title, subtitle, source, output)
        print(output)


if __name__ == "__main__":
    main()
